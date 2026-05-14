import os
import re
from typing import Any, Dict, List, Optional


def _clean_text(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _is_heading_style(style_name: Optional[str]) -> bool:
    if not style_name:
        return False
    name = style_name.strip().lower()
    if name.startswith("heading"):
        return True
    if name.startswith("titre"):
        return True
    if name.startswith("title"):
        return True
    return False


def parse_docx(
    file_path: str,
    *,
    include_tables: bool = True,
) -> Dict[str, Any]:
    if not os.path.exists(file_path):
        raise FileNotFoundError(file_path)

    from docx import Document

    doc = Document(file_path)

    items: List[Dict[str, Any]] = []
    current_section: Optional[str] = None

    paragraph_index = 0
    for p in doc.paragraphs:
        paragraph_index += 1
        raw = p.text or ""
        text = _clean_text(raw)
        if not text:
            continue

        style_name = getattr(getattr(p, "style", None), "name", None)
        is_heading = _is_heading_style(style_name)
        if is_heading:
            current_section = text

        items.append(
            {
                "type": "paragraph",
                "index": paragraph_index,
                "style": style_name,
                "is_heading": is_heading,
                "section": current_section,
                "text": text,
            }
        )

    if include_tables:
        table_index = 0
        for t in doc.tables:
            table_index += 1
            rows: List[List[str]] = []
            for row in t.rows:
                row_texts: List[str] = []
                for cell in row.cells:
                    cell_text = _clean_text(cell.text or "")
                    row_texts.append(cell_text)
                if any(r for r in row_texts):
                    rows.append(row_texts)
            if not rows:
                continue
            items.append(
                {
                    "type": "table",
                    "index": table_index,
                    "section": current_section,
                    "rows": rows,
                }
            )

    return {
        "file_path": file_path,
        "file_name": os.path.basename(file_path),
        "items": items,
    }

