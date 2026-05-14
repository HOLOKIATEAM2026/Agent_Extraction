import os
import re
from typing import Any, Dict, List, Optional

from docx import Document


def _clean_text(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _is_heading_style(style_name: Optional[str]) -> bool:
    if not style_name:
        return False
    return style_name.lower().startswith("heading")


def parse_docx(file_path: str) -> Dict[str, Any]:
    if not os.path.exists(file_path):
        raise FileNotFoundError(file_path)

    doc = Document(file_path)

    blocks: List[Dict[str, Any]] = []
    current_section: Optional[str] = None
    idx = 0

    for p in doc.paragraphs:
        text = _clean_text(p.text or "")
        style_name = getattr(getattr(p, "style", None), "name", None)
        is_heading = _is_heading_style(style_name)

        if is_heading and text:
            current_section = text

        if text:
            blocks.append(
                {
                    "index": idx,
                    "type": "paragraph",
                    "text": text,
                    "style": style_name,
                    "is_heading": is_heading,
                    "section": current_section,
                }
            )
            idx += 1

    for t_i, table in enumerate(doc.tables):
        rows: List[List[str]] = []
        for row in table.rows:
            cells = [_clean_text(c.text or "") for c in row.cells]
            rows.append(cells)

        blocks.append(
            {
                "index": idx,
                "type": "table",
                "table_index": t_i,
                "rows": rows,
                "section": current_section,
            }
        )
        idx += 1

    return {
        "file_path": file_path,
        "file_name": os.path.basename(file_path),
        "blocks": blocks,
        "block_count": len(blocks),
    }

